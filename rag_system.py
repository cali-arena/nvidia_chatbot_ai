"""
Advanced RAG System with Large File Support (up to 2GB)
Includes: Vector Store, Embeddings, RAG Agents, and Evaluation
"""

import os
import asyncio
import tempfile
import hashlib
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
import logging
import warnings

# Suprimir warnings e logs desnecessários
warnings.filterwarnings("ignore")
logging.getLogger("sentence_transformers").setLevel(logging.ERROR)
logging.getLogger("transformers").setLevel(logging.ERROR)
logging.getLogger("torch").setLevel(logging.ERROR)
logging.getLogger("chromadb").setLevel(logging.ERROR)
logging.getLogger("langchain").setLevel(logging.ERROR)

# Core imports
import streamlit as st
import pandas as pd
import numpy as np
from pathlib import Path

# LangChain imports
from langchain.text_splitter import RecursiveCharacterTextSplitter, TokenTextSplitter
from langchain.docstore.document import Document
try:
    from langchain_huggingface import HuggingFaceEmbeddings
except ImportError:
    from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma, FAISS
from langchain.memory import ConversationBufferMemory
from langchain.schema import BaseRetriever

# NVIDIA AI imports
from langchain_nvidia_ai_endpoints import ChatNVIDIA, NVIDIAEmbeddings

# Document processing
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
try:
    import tiktoken
except ImportError:
    tiktoken = None

# Evaluation
try:
    from sklearn.metrics.pairwise import cosine_similarity
    SKLEARN_AVAILABLE = True
except ImportError:
    SKLEARN_AVAILABLE = False
    cosine_similarity = None

try:
    import matplotlib.pyplot as plt
    import seaborn as sns
    PLOTTING_AVAILABLE = True
except ImportError:
    PLOTTING_AVAILABLE = False
    plt = None
    sns = None

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class RAGConfig:
    """Configuration for RAG system"""
    chunk_size: int = 1000
    chunk_overlap: int = 200
    max_file_size_gb: float = 2.0
    embedding_model: str = "nvidia/Llama-3.2-3B-Instruct-TensorRT-LLM"
    vector_store_type: str = "chroma"  # chroma or faiss
    retrieval_k: int = 5
    similarity_threshold: float = 0.7

class LargeFileProcessor:
    """Handles large file processing with chunking strategies"""
    
    def __init__(self, config: RAGConfig):
        self.config = config
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.chunk_size,
            chunk_overlap=config.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
    def get_file_hash(self, file_path: str) -> str:
        """Generate hash for file to avoid reprocessing"""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()
    
    def validate_file_size(self, file_path: str) -> bool:
        """Validate file size is under limit"""
        size_gb = os.path.getsize(file_path) / (1024**3)
        return size_gb <= self.config.max_file_size_gb
    
    async def process_large_pdf(self, file_path: str) -> List[Document]:
        """Process large PDF files with streaming and memory optimization"""
        documents = []
        
        try:
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                total_pages = len(reader.pages)
                
                # Process pages in batches to manage memory
                batch_size = 50  # Process 50 pages at a time
                
                for batch_start in range(0, total_pages, batch_size):
                    batch_end = min(batch_start + batch_size, total_pages)
                    
                    for page_num in range(batch_start, batch_end):
                        try:
                            page = reader.pages[page_num]
                            text = page.extract_text()
                            
                            if text.strip():
                                # Split large pages into smaller chunks if needed
                                if len(text) > self.config.chunk_size * 2:
                                    chunks = self.text_splitter.split_text(text)
                                    for chunk_idx, chunk in enumerate(chunks):
                                        doc = Document(
                                            page_content=chunk,
                                            metadata={
                                                "source": file_path,
                                                "page": page_num + 1,
                                                "chunk": chunk_idx,
                                                "type": "pdf",
                                                "file_hash": self.get_file_hash(file_path),
                                                "total_chunks": len(chunks)
                                            }
                                        )
                                        documents.append(doc)
                                else:
                                    doc = Document(
                                        page_content=text,
                                        metadata={
                                            "source": file_path,
                                            "page": page_num + 1,
                                            "type": "pdf",
                                            "file_hash": self.get_file_hash(file_path)
                                        }
                                    )
                                    documents.append(doc)
                                    
                        except Exception as e:
                            logger.warning(f"Error processing page {page_num}: {e}")
                            continue
                    
                    # Log progress for large files
                    if total_pages > 100:
                        logger.info(f"Processed pages {batch_start}-{batch_end} of {total_pages}")
                        
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            
        return documents
    
    async def process_large_docx(self, file_path: str) -> List[Document]:
        """Process large DOCX files with memory optimization"""
        documents = []
        
        try:
            doc = DocxDocument(file_path)
            full_text = []
            
            # Process paragraphs in batches to manage memory
            paragraphs = doc.paragraphs
            batch_size = 1000  # Process 1000 paragraphs at a time
            
            for batch_start in range(0, len(paragraphs), batch_size):
                batch_end = min(batch_start + batch_size, len(paragraphs))
                batch_paragraphs = paragraphs[batch_start:batch_end]
                
                for paragraph in batch_paragraphs:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text)
                
                # Log progress for large files
                if len(paragraphs) > 5000:
                    logger.info(f"Processed paragraphs {batch_start}-{batch_end} of {len(paragraphs)}")
            
            # Split into chunks
            text = "\n".join(full_text)
            chunks = self.text_splitter.split_text(text)
            
            for i, chunk in enumerate(chunks):
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "source": file_path,
                        "chunk": i,
                        "type": "docx",
                        "file_hash": self.get_file_hash(file_path),
                        "total_chunks": len(chunks)
                    }
                )
                documents.append(doc)
                
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            
        return documents
    
    async def process_file(self, file_path: str, file_type: str) -> List[Document]:
        """Process any supported file type with optimized memory usage"""
        if not self.validate_file_size(file_path):
            raise ValueError(f"File size exceeds {self.config.max_file_size_gb}GB limit")
        
        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
        logger.info(f"Processing {file_type} file: {file_path} ({file_size_mb:.2f} MB)")
        
        if file_type.lower() == "pdf":
            return await self.process_large_pdf(file_path)
        elif file_type.lower() in ["docx", "doc"]:
            return await self.process_large_docx(file_path)
        else:
            # Handle text files with streaming for large files
            documents = []
            
            if file_size_mb > 100:  # For files larger than 100MB, use streaming
                logger.info("Using streaming processing for large text file")
                with open(file_path, 'r', encoding='utf-8') as f:
                    buffer = ""
                    chunk_count = 0
                    
                    for line in f:
                        buffer += line
                        
                        # Process buffer when it reaches chunk_size
                        if len(buffer) >= self.config.chunk_size:
                            chunks = self.text_splitter.split_text(buffer)
                            
                            for chunk in chunks[:-1]:  # Keep last chunk in buffer
                                doc = Document(
                                    page_content=chunk,
                                    metadata={
                                        "source": file_path,
                                        "chunk": chunk_count,
                                        "type": file_type,
                                        "file_hash": self.get_file_hash(file_path)
                                    }
                                )
                                documents.append(doc)
                                chunk_count += 1
                            
                            # Keep the last chunk in buffer
                            buffer = chunks[-1] if chunks else ""
                    
                    # Process remaining buffer
                    if buffer.strip():
                        chunks = self.text_splitter.split_text(buffer)
                        for chunk in chunks:
                            doc = Document(
                                page_content=chunk,
                                metadata={
                                    "source": file_path,
                                    "chunk": chunk_count,
                                    "type": file_type,
                                    "file_hash": self.get_file_hash(file_path)
                                }
                            )
                            documents.append(doc)
                            chunk_count += 1
            else:
                # For smaller files, use regular processing
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                
                chunks = self.text_splitter.split_text(text)
                
                for i, chunk in enumerate(chunks):
                    doc = Document(
                        page_content=chunk,
                        metadata={
                            "source": file_path,
                            "chunk": i,
                            "type": file_type,
                            "file_hash": self.get_file_hash(file_path)
                        }
                    )
                    documents.append(doc)
            
            return documents

class VectorStoreManager:
    """Manages vector stores and embeddings"""
    
    def __init__(self, config: RAGConfig, api_key: str):
        self.config = config
        self.api_key = api_key
        self.embeddings = None
        self.vector_store = None
        self._initialize_embeddings()
    
    def _initialize_embeddings(self):
        """Initialize embedding model"""
        try:
            # Try to use HuggingFace embeddings with proper device handling
            import torch
            device = "cpu"  # Force CPU to avoid meta tensor issues
            
            self.embeddings = HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-MiniLM-L6-v2",
                model_kwargs={'device': device}
            )
            logger.info("Using HuggingFace embeddings with CPU device")
        except Exception as e:
            logger.error(f"Failed to initialize HuggingFace embeddings: {e}")
            # Fallback to simple embeddings
            from langchain.embeddings import FakeEmbeddings
            self.embeddings = FakeEmbeddings(size=384)
            logger.info("Using fake embeddings as fallback")
    
    def create_vector_store(self, documents: List[Document], collection_name: str = "default"):
        """Create vector store from documents"""
        if not documents:
            raise ValueError("No documents provided")
        
        try:
            if self.config.vector_store_type == "chroma":
                self.vector_store = Chroma.from_documents(
                    documents=documents,
                    embedding=self.embeddings,
                    collection_name=collection_name,
                    persist_directory="./chroma_db"
                )
            else:  # FAISS
                self.vector_store = FAISS.from_documents(
                    documents=documents,
                    embedding=self.embeddings
                )
            
            logger.info(f"Created {self.config.vector_store_type} vector store with {len(documents)} documents")
            
        except Exception as e:
            logger.error(f"Error creating vector store: {e}")
            raise
    
    def add_documents(self, documents: List[Document]):
        """Add documents to existing vector store"""
        if not self.vector_store:
            raise ValueError("Vector store not initialized")
        
        self.vector_store.add_documents(documents)
        logger.info(f"Added {len(documents)} documents to vector store")
    
    def similarity_search(self, query: str, k: int = None) -> List[Document]:
        """Perform similarity search"""
        if not self.vector_store:
            raise ValueError("Vector store not initialized")
        
        k = k or self.config.retrieval_k
        results = self.vector_store.similarity_search(query, k=k)
        return results
    
    def similarity_search_with_score(self, query: str, k: int = None) -> List[Tuple[Document, float]]:
        """Perform similarity search with scores"""
        if not self.vector_store:
            raise ValueError("Vector store not initialized")
        
        k = k or self.config.retrieval_k
        
        if hasattr(self.vector_store, 'similarity_search_with_score'):
            return self.vector_store.similarity_search_with_score(query, k=k)
        else:
            # Fallback for stores without score method
            docs = self.similarity_search(query, k)
            return [(doc, 1.0) for doc in docs]  # Dummy scores

class RAGAgent:
    """RAG Agent with different strategies"""
    
    def __init__(self, vector_store_manager: VectorStoreManager, config: RAGConfig, api_key: str):
        self.vector_manager = vector_store_manager
        self.config = config
        self.api_key = api_key
        self.llm = None
        self.qa_chain = None
        self._initialize_llm()
    
    def _initialize_llm(self):
        """Initialize LLM with optimal parameters"""
        self.llm = ChatNVIDIA(
            model="meta/llama-3.1-405b-instruct",
            temperature=0.3,
            max_completion_tokens=1000,
            nvidia_api_key=self.api_key
        )
    
    def create_retrieval_chain(self, memory: bool = True):
        """Create retrieval QA chain"""
        if not self.vector_manager.vector_store:
            raise ValueError("Vector store not initialized")
        
        retriever = self.vector_manager.vector_store.as_retriever(
            search_kwargs={"k": self.config.retrieval_k}
        )
        
        # Simplified chain creation without deprecated imports
        self.retriever = retriever
        if memory:
            self.memory = ConversationBufferMemory(
                memory_key="chat_history",
                return_messages=True
            )
        else:
            self.memory = None
    
    def query(self, question: str, use_memory: bool = True) -> Dict[str, Any]:
        """Query the RAG system"""
        try:
            # Get relevant documents
            relevant_docs = self.vector_manager.similarity_search_with_score(question)
            
            # Filter by similarity threshold
            filtered_docs = [
                doc for doc, score in relevant_docs 
                if score >= self.config.similarity_threshold
            ]
            
            if not filtered_docs:
                return {
                    "answer": "No relevant information found in the documents.",
                    "source_documents": [],
                    "confidence": 0.0
                }
            
            # Create context from relevant documents
            context = "\n\n".join([doc.page_content for doc in filtered_docs])
            
            # Create prompt with context
            prompt = f"""Based on the following context from documents, answer the question: {question}

Context:
{context}

Please provide a comprehensive answer based on the provided context. If the context doesn't contain enough information, say so clearly."""

            # Generate answer using LLM directly
            response = self.llm.invoke(prompt)
            answer = response.content
            
            # Calculate confidence based on document scores
            avg_score = np.mean([score for _, score in relevant_docs])
            confidence = min(avg_score, 1.0)
            
            return {
                "answer": answer,
                "source_documents": filtered_docs,
                "confidence": confidence,
                "raw_response": response
            }
            
        except Exception as e:
            logger.error(f"Error in RAG query: {e}")
            return {
                "answer": f"Error processing query: {str(e)}",
                "source_documents": [],
                "confidence": 0.0
            }

class RAGEvaluator:
    """Evaluates RAG system performance"""
    
    def __init__(self):
        self.evaluation_results = []
    
    def evaluate_response(self, question: str, answer: str, 
                         source_docs: List[Document], 
                         ground_truth: str = None) -> Dict[str, Any]:
        """Evaluate a single RAG response"""
        
        evaluation = {
            "question": question,
            "answer": answer,
            "num_source_docs": len(source_docs),
            "avg_doc_length": np.mean([len(doc.page_content) for doc in source_docs]) if source_docs else 0,
            "timestamp": pd.Timestamp.now()
        }
        
        # Relevance score (simple heuristic)
        if source_docs:
            relevance_score = min(len(source_docs) / 5, 1.0)  # Normalize to 0-1
        else:
            relevance_score = 0.0
        
        evaluation["relevance_score"] = relevance_score
        
        # Answer completeness (simple heuristic)
        answer_length = len(answer.split())
        completeness_score = min(answer_length / 50, 1.0)  # Normalize to 0-1
        evaluation["completeness_score"] = completeness_score
        
        # Overall score
        evaluation["overall_score"] = (relevance_score + completeness_score) / 2
        
        self.evaluation_results.append(evaluation)
        return evaluation
    
    def get_evaluation_summary(self) -> Dict[str, Any]:
        """Get summary of all evaluations"""
        if not self.evaluation_results:
            return {"message": "No evaluations available"}
        
        df = pd.DataFrame(self.evaluation_results)
        
        summary = {
            "total_evaluations": len(df),
            "avg_relevance_score": df["relevance_score"].mean(),
            "avg_completeness_score": df["completeness_score"].mean(),
            "avg_overall_score": df["overall_score"].mean(),
            "avg_source_docs": df["num_source_docs"].mean(),
            "avg_doc_length": df["avg_doc_length"].mean()
        }
        
        return summary
    
    def plot_evaluation_metrics(self):
        """Create visualization of evaluation metrics"""
        if not self.evaluation_results:
            return None
        
        df = pd.DataFrame(self.evaluation_results)
        
        fig, axes = plt.subplots(2, 2, figsize=(12, 8))
        
        # Relevance scores over time
        axes[0, 0].plot(df.index, df["relevance_score"])
        axes[0, 0].set_title("Relevance Scores Over Time")
        axes[0, 0].set_ylabel("Relevance Score")
        
        # Completeness scores over time
        axes[0, 1].plot(df.index, df["completeness_score"])
        axes[0, 1].set_title("Completeness Scores Over Time")
        axes[0, 1].set_ylabel("Completeness Score")
        
        # Distribution of overall scores
        axes[1, 0].hist(df["overall_score"], bins=10, alpha=0.7)
        axes[1, 0].set_title("Distribution of Overall Scores")
        axes[1, 0].set_xlabel("Overall Score")
        axes[1, 0].set_ylabel("Frequency")
        
        # Source documents vs scores
        axes[1, 1].scatter(df["num_source_docs"], df["overall_score"])
        axes[1, 1].set_title("Source Documents vs Overall Score")
        axes[1, 1].set_xlabel("Number of Source Documents")
        axes[1, 1].set_ylabel("Overall Score")
        
        plt.tight_layout()
        return fig

class AdvancedRAGSystem:
    """Main RAG system orchestrator"""
    
    def __init__(self, api_key: str, config: RAGConfig = None):
        self.config = config or RAGConfig()
        self.api_key = api_key
        
        # Initialize components
        self.file_processor = LargeFileProcessor(self.config)
        self.vector_manager = VectorStoreManager(self.config, api_key)
        self.rag_agent = RAGAgent(self.vector_manager, self.config, api_key)
        self.evaluator = RAGEvaluator()
        
        # Session state
        self.processed_files = set()
        self.conversation_memory = ConversationBufferMemory(
            memory_key="chat_history",
            return_messages=True
        )
    
    async def process_file(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """Process a file and add to vector store"""
        try:
            # Check if already processed
            file_hash = self.file_processor.get_file_hash(file_path)
            if file_hash in self.processed_files:
                return {"status": "already_processed", "message": "File already processed"}
            
            # Process file
            documents = await self.file_processor.process_file(file_path, file_type)
            
            if not documents:
                return {"status": "error", "message": "No content extracted from file"}
            
            # Add to vector store
            if not self.vector_manager.vector_store:
                self.vector_manager.create_vector_store(documents)
            else:
                self.vector_manager.add_documents(documents)
            
            # Mark as processed
            self.processed_files.add(file_hash)
            
            return {
                "status": "success",
                "message": f"Processed {len(documents)} chunks",
                "chunks": len(documents),
                "file_hash": file_hash
            }
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            return {"status": "error", "message": str(e)}
    
    def query(self, question: str, use_memory: bool = True) -> Dict[str, Any]:
        """Query the RAG system"""
        try:
            result = self.rag_agent.query(question, use_memory)
            
            # Evaluate response
            evaluation = self.evaluator.evaluate_response(
                question=question,
                answer=result["answer"],
                source_docs=result["source_documents"]
            )
            
            result["evaluation"] = evaluation
            return result
            
        except Exception as e:
            logger.error(f"Error in RAG query: {e}")
            return {
                "answer": f"Error: {str(e)}",
                "source_documents": [],
                "confidence": 0.0,
                "evaluation": {"overall_score": 0.0}
            }
    
    def get_system_stats(self) -> Dict[str, Any]:
        """Get system statistics"""
        stats = {
            "processed_files": len(self.processed_files),
            "vector_store_type": self.config.vector_store_type,
            "chunk_size": self.config.chunk_size,
            "max_file_size_gb": self.config.max_file_size_gb,
            "embedding_model": self.config.embedding_model
        }
        
        if self.vector_manager.vector_store:
            try:
                # Get collection info if available
                if hasattr(self.vector_manager.vector_store, '_collection'):
                    collection = self.vector_manager.vector_store._collection
                    stats["total_documents"] = collection.count()
                else:
                    stats["total_documents"] = "Unknown"
            except:
                stats["total_documents"] = "Unknown"
        else:
            stats["total_documents"] = 0
        
        # Add evaluation summary
        eval_summary = self.evaluator.get_evaluation_summary()
        stats["evaluation_summary"] = eval_summary
        
        return stats
